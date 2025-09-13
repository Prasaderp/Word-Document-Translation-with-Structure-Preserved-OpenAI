import asyncio
import re
import random
from typing import List, Dict, Tuple
from openai import AsyncOpenAI
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from prompts import (
    SYSTEM_PROMPT_BASE,
    USER_TERMS_INSTRUCTION,
    MASK_INSTRUCTION,
    RETRY_PROMPT_ADDITION,
    QUALITY_ASSESSMENT_PROMPT
)

class EnhancedTranslator:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.client = AsyncOpenAI(api_key=api_key)
        self.quality_threshold = 30
        self.max_retries = 3
        self.concurrency_limit = 10
        self.rpm_limit = 950
        self._spacy_nlp = None
        self._spacy_available = None

    def ensure_spacy(self) -> bool:
        if self._spacy_available is not None:
            return self._spacy_available
        try:
            import spacy
            try:
                self._spacy_nlp = spacy.load("en_core_web_sm")
                self._spacy_available = True
            except Exception:
                self._spacy_nlp = None
                self._spacy_available = False
        except Exception:
            self._spacy_nlp = None
            self._spacy_available = False
        return self._spacy_available

    def _find_user_term_spans(self, text: str, user_terms: List[str]) -> List[Tuple[int, int, str]]:
        if not user_terms or not text:
            return []
        spans = []
        occupied = [False] * len(text)
        sorted_terms = sorted({t for t in user_terms if t}, key=lambda x: len(x), reverse=True)
        for term in sorted_terms:
            escaped = re.escape(term)
            pattern = re.compile(rf"(?i)(?<!\w){escaped}(?!\w)")
            matches = list(pattern.finditer(text))
            if not matches:
                fallback = re.compile(rf"(?i){escaped}")
                matches = list(fallback.finditer(text))
            for m in matches:
                s, e = m.start(), m.end()
                if any(occupied[i] for i in range(s, e)):
                    continue
                spans.append((s, e, text[s:e]))
                for i in range(s, e):
                    occupied[i] = True
        spans.sort(key=lambda x: x[0])
        return spans

    def _find_spacy_entity_spans(self, text: str) -> List[Tuple[int, int, str]]:
        if not text or not self.ensure_spacy() or not self._spacy_nlp:
            return []
        doc = self._spacy_nlp(text)
        spans = []
        allowed_labels = {
            "PERSON", "ORG", "GPE", "LOC", "FAC", "NORP", "PRODUCT", "EVENT", "WORK_OF_ART", "LAW", "LANGUAGE"
        }
        for ent in getattr(doc, "ents", []):
            if getattr(ent, "label_", None) not in allowed_labels:
                continue
            s, e = int(ent.start_char), int(ent.end_char)
            if 0 <= s < e <= len(text):
                spans.append((s, e, text[s:e]))
        spans.sort(key=lambda x: x[0])
        return spans

    def _mask_text(self, text: str, user_terms: List[str]) -> Tuple[str, Dict[str, str]]:
        if not text:
            return text, {}
        user_spans = self._find_user_term_spans(text, user_terms)
        spacy_spans = self._find_spacy_entity_spans(text)
        filtered_spacy_spans = []
        for s_s, s_e, s_val in spacy_spans:
            if not any(not (s_e <= u_s or s_s >= u_e) for u_s, u_e, _ in user_spans):
                filtered_spacy_spans.append((s_s, s_e, s_val))
        combined = [(s, e, v, "UT") for s, e, v in user_spans] + [(s, e, v, "NE") for s, e, v in filtered_spacy_spans]
        if not combined:
            return text, {}
        combined.sort(key=lambda x: x[0])
        parts, token_map, cursor, ut_idx, ne_idx = [], {}, 0, 0, 0
        for s, e, v, kind in combined:
            if s < cursor: continue
            parts.append(text[cursor:s])
            token = f"<<UT{ut_idx}>>" if kind == "UT" else f"<<NE{ne_idx}>>"
            if kind == "UT": ut_idx += 1
            else: ne_idx += 1
            parts.append(token)
            token_map[token] = v
            cursor = e
        parts.append(text[cursor:])
        return "".join(parts), token_map

    def _unmask_text(self, text: str, token_map: Dict[str, str]) -> str:
        if not token_map or not text:
            return text
        for token, original in sorted(token_map.items(), key=lambda x: -len(x[0])):
            text = text.replace(token, original)
        return text

    def is_translatable(self, text: str) -> bool:
        return bool(text and text.strip() and any(char.isalpha() for char in text))

    def copy_run_style(self, source_run, target_run):
        target_run.style = source_run.style
        target_run.bold, target_run.italic, target_run.underline = source_run.bold, source_run.italic, source_run.underline
        font, source_font = target_run.font, source_run.font
        font.name, font.size = source_font.name, source_font.size
        if source_font.color and source_font.color.rgb:
            font.color.rgb = source_font.color.rgb

    def get_all_paragraphs(self, doc: Document) -> List[Paragraph]:
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
        for section in doc.sections:
            all_paragraphs.extend(section.header.paragraphs)
            all_paragraphs.extend(section.footer.paragraphs)
        return all_paragraphs

    async def translate_text_with_quality(self, text: str, target_lang: str, user_terms: List[str]) -> tuple:
        base_delay = 1.0
        for attempt in range(self.max_retries):
            try:
                masked_text, token_map = self._mask_text(text, user_terms or [])
                system_prompt = SYSTEM_PROMPT_BASE.format(target_lang=target_lang)
                if user_terms:
                    terms_list_str = ", ".join(f'"{term}"' for term in user_terms)
                    system_prompt += USER_TERMS_INSTRUCTION.format(terms_list_str=terms_list_str)
                system_prompt += MASK_INSTRUCTION
                if attempt > 0:
                    system_prompt += RETRY_PROMPT_ADDITION.format(attempt=attempt + 1)
                
                message = [{"role": "system", "content": system_prompt}, {"role": "user", "content": masked_text}]
                response = await self.client.chat.completions.create(
                    model="gpt-4o-mini", messages=message,
                    temperature=0.1 if attempt > 0 else 0, max_tokens=4000
                )
                translated_text = self._unmask_text(response.choices[0].message.content.strip(), token_map)
                
                if not translated_text:
                    if attempt == self.max_retries - 1: return text, 0
                    continue
                
                quality_score = await self.validate_translation_quality(text, translated_text, target_lang)
                if quality_score >= self.quality_threshold or attempt == self.max_retries - 1:
                    return translated_text, quality_score
                
                print(f"[INFO] Quality score {quality_score} below threshold {self.quality_threshold}, retrying...")
            except Exception as e:
                print(f"[WARNING] Translation attempt {attempt + 1} failed: {e}")
                if attempt == self.max_retries - 1: return text, 0
            
            await asyncio.sleep(base_delay * (2 ** attempt) + random.uniform(0, 1))
        return text, 0

    async def validate_translation_quality(self, original: str, translated: str, target_lang: str) -> int:
        try:
            prompt = QUALITY_ASSESSMENT_PROMPT.format(target_lang=target_lang, original=original, translated=translated)
            response = await self.client.chat.completions.create(
                model="gpt-4o", messages=[{"role": "user", "content": prompt}],
                temperature=0, max_tokens=10
            )
            score_text = response.choices[0].message.content.strip()
            score = int(re.search(r'\d+', score_text).group())
            return max(0, min(40, score))
        except Exception:
            return 20

    async def process_enhanced_translation(self, source_path: str, output_path: str, target_language: str, user_terms: List[str] = None):
        doc = docx.Document(source_path)
        all_paragraphs = self.get_all_paragraphs(doc)
        prefix_pattern = re.compile(r'^\s*(?:\d+\.\s*|\(\d+\)\s*|[a-zA-Z]\.\s*|\([a-zA-Z]\)\s*)')
        
        texts_for_translation = {}
        for p in all_paragraphs:
            core_text = p.text
            match = prefix_pattern.match(core_text)
            prefix = match.group(0) if match else ""
            if match: core_text = core_text[len(prefix):]
            if self.is_translatable(core_text):
                texts_for_translation[core_text] = {"original_text": p.text, "prefix": prefix, "paragraph": p}
        
        unique_texts_to_translate = list(texts_for_translation.keys())
        print(f"[INFO] Enhanced translation: Found {len(unique_texts_to_translate)} unique text segments")
        
        translated_cache, quality_scores = {}, []
        semaphore = asyncio.Semaphore(self.concurrency_limit)
        request_delay = 60.0 / self.rpm_limit
        
        async def translate_task(text):
            async with semaphore:
                translated_text, quality_score = await self.translate_text_with_quality(text, target_language, user_terms or [])
                await asyncio.sleep(request_delay)
                return text, translated_text, quality_score

        if unique_texts_to_translate:
            total_texts = len(unique_texts_to_translate)
            print(f"[INFO] Starting enhanced translation to {target_language} with quality control...")
            tasks = [translate_task(text) for text in unique_texts_to_translate]
            for i, task in enumerate(asyncio.as_completed(tasks)):
                original_text, translated_text, quality_score = await task
                translated_cache[original_text] = translated_text
                quality_scores.append(quality_score)
                progress = (i + 1) / total_texts * 100
                avg_quality = sum(quality_scores) / len(quality_scores) if quality_scores else 0
                yield progress, avg_quality
            
            avg_quality_final = sum(quality_scores) / len(quality_scores) if quality_scores else 0
            print(f"\n[INFO] Enhanced translation complete. Average quality: {avg_quality_final:.1f}/40")
        
        print("[INFO] Applying enhanced translations to document...")
        for para in all_paragraphs:
            original_text = para.text
            match = prefix_pattern.match(original_text)
            prefix = match.group(0) if match else ""
            core_text = original_text[len(prefix):] if match else original_text
            
            if core_text in translated_cache:
                translated_core_text = translated_cache[core_text]
                final_text = prefix + translated_core_text
                first_run = para.runs[0] if para.runs else None
                para.clear()
                new_run = para.add_run(final_text)
                if first_run:
                    self.copy_run_style(first_run, new_run)
        
        doc.save(output_path)
        print(f"[SUCCESS] Enhanced translated document saved to {output_path}")